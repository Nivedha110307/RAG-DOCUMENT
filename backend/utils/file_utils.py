"""
file_utils.py — File I/O and text extraction utilities.

WHY SEPARATE EXTRACTORS?
- PDF: Binary format, needs PyMuPDF or pdfplumber (layout-aware)
- DOCX: ZIP archive with XML, needs python-docx
- TXT: Direct read
- Each has quirks (PDF OCR, DOCX tables, TXT encoding)
"""

import logging
from pathlib import Path
from typing import Optional

import aiofiles
from fastapi import UploadFile

logger = logging.getLogger(__name__)


async def save_upload(file: UploadFile, upload_dir: str) -> Path:
    """Save uploaded file to disk, return the saved path."""
    dest = Path(upload_dir)
    dest.mkdir(parents=True, exist_ok=True)
    
    file_path = dest / file.filename
    
    async with aiofiles.open(file_path, "wb") as f:
        content = await file.read()
        await f.write(content)
    
    return file_path


def extract_text_from_pdf(file_path: Path) -> tuple[str, Optional[list[dict]]]:
    """
    Extract text from PDF using PyMuPDF (fitz).
    
    Why PyMuPDF?
    - Handles complex layouts better than PyPDF2
    - Extracts text in reading order
    - Returns page-level metadata
    - Supports encrypted PDFs (with password)
    
    Returns: (full_text, page_metadata_per_page)
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        raise ImportError("PyMuPDF not installed. Run: pip install pymupdf")
    
    doc = fitz.open(str(file_path))
    pages_text = []
    page_metadata = []
    
    for page_num, page in enumerate(doc, 1):
        text = page.get_text("text")  # "text" = plain, "html" = with formatting
        if text.strip():
            pages_text.append(text)
            page_metadata.append({"page_number": page_num})
        logger.debug("Page %d: %d chars", page_num, len(text))
    
    doc.close()
    
    full_text = "\n\n".join(pages_text)
    return full_text, page_metadata


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from Word documents using python-docx.
    
    Handles:
    - Paragraphs and headings
    - Tables (flattened to text)
    - Headers/footers (skipped to avoid repetition)
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    doc = Document(str(file_path))
    parts = []
    
    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    
    # Extract tables (flatten to readable format)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)
    
    return "\n\n".join(parts)


def compute_file_hash(file_path: Path) -> str:
    """SHA256 hash of file content — useful for deduplication."""
    import hashlib
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            sha256.update(chunk)
    return sha256.hexdigest()
